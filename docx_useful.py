from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement

'''
1. Add caption to table or figure

To make an intiutive document, it necessary to caption figures and tables properly. The docx documentation doen not direcly tell you hoe to get this done.

Inputs:
text: caption text
doc_instance: the document instance into which the caption needs to be added
caption_type: 1 for Table caption, 2 for Figure Caption. Anything else will throw an error
'''

def add_captions(text, doc_instance, caption_type):
    try:
        if caption_type not in [1,2]:
            raise ValueError ("the input for caption_type should be 1 for table caption or 2 for figure caption")

        if caption_type == 1:
            para = doc_instance.add_paragraph("Table ")
        if caption_type == 2:
            para = doc_instance.add_paragraph("Figure ")

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER #To align the captoion to center

        run = para.add_run()
        r = run._r
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)

        instrText = OxmlElement('w:instrText')

        if caption_type==1:
            instrText.text = ' SEQ Table \\* ARABIC'
        if caption_type == 2:
            instrText.text = ' SEQ Figure \\* ARABIC'
        
        r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)

        para.add_run(f": {text}")
    
    except Exception as e:
        print(e)



'''
2. Include List of Tables or Figures at the beginning of the document

If the figures and tables are captioned properly as using the functionality shown above, 
then the defined sequence of tables and figures can be used to make the list of tables / figures 

Inputs:
run: the run object in the document paragraph insatnce
list_type: 1 for Table caption, 2 for Figure Caption. Anything else will throw an error
'''

def list_tabs_figs(run_obj, list_type):
    try:
        if list_type not in [1,2]:
            raise ValueError ("the input for list_type should be 1 for list of tables or 2 for list of figures")
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        fldChar.set(qn('w:dirty'), 'true')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')

        if list_type == 1:
            instrText.text = 'TOC \\h \\z \\c "Table"'
        if list_type == 2:
            instrText.text = 'TOC \\h \\z \\c "Figure"'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        run_obj._r.append(fldChar)
        run_obj._r.append(instrText)
        run_obj._r.append(fldChar2)
        run_obj._r.append(fldChar4)
    
    except Exception as e:
        print(e)





